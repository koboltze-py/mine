"""
Report Generator für PDF, Word, ODF und Apple Pages Dokumente
"""
import os
import zipfile
import xml.etree.ElementTree as ET
import re
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    def __init__(self, output_dir: str = "reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    # Font-Mapping: benutzerfreundliche Namen -> ReportLab-Namen
    _PDF_FONTS = {
        "Arial": "Helvetica",
        "Times New Roman": "Times-Roman",
        "Courier New": "Courier",
    }
    _DEFAULT_FONT = "Helvetica"
    _DEFAULT_SIZE = 11

    # Vitalwerte Labels und Einheiten
    _VITAL_LABELS = [
        ('rr',    'RR',     'mmHg'),
        ('hf',    'HF',     '/min'),
        ('spo2',  'SpO₂',   '%'),
        ('spco',  'SpCO',   '%'),
        ('af',    'AF',     '/min'),
        ('bz',    'BZ',     'mmol/l'),
        ('temp',  'Temp',   '°C'),
        ('gcs',   'GCS',    ''),
        ('etco2', 'EtCO₂',  'mmHg'),
    ]

    def generate_pdf(self, titel: str, thema: str, inhalt: str, bericht_id: int,
                     font_family: str = "Arial", font_size: int = 11,
                     reflexion: str = "", abcde_data: dict = None, vitalwerte: dict = None) -> str:
        """
        Generiert ein PDF-Dokument des Einsatzberichts
        
        Args:
            titel: Titel des Berichts
            thema: Thema des Berichts
            inhalt: Inhalt des Berichts
            bericht_id: ID des Berichts
        
        Returns:
            Pfad zur erstellten PDF-Datei
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"bericht_{bericht_id}_{timestamp}.pdf"
        filepath = os.path.join(self.output_dir, filename)

        rl_font = self._PDF_FONTS.get(font_family, self._DEFAULT_FONT)
        fs = font_size or self._DEFAULT_SIZE

        # PDF erstellen
        doc = SimpleDocTemplate(filepath, pagesize=A4,
                              rightMargin=2*cm, leftMargin=2*cm,
                              topMargin=2*cm, bottomMargin=2*cm)

        # Styles
        styles = getSampleStyleSheet()

        # Custom Styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=rl_font + '-Bold' if rl_font == 'Helvetica' else rl_font,
            fontSize=fs + 13,
            textColor='#1a1a1a',
            spaceAfter=30,
            alignment=1  # Center
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=rl_font + '-Bold' if rl_font == 'Helvetica' else rl_font,
            fontSize=fs + 3,
            textColor='#333333',
            spaceAfter=6,
            spaceBefore=10,
            keepWithNext=1
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontName=rl_font,
            fontSize=fs,
            textColor='#000000',
            spaceAfter=12,
            alignment=4  # Justify
        )
        
        # Story (Inhalt)
        story = []
        
        # Titel
        story.append(Paragraph("EINSATZBERICHT", title_style))
        story.append(Spacer(1, 0.5*cm))
        
        # Titel und Alarmierung
        story.append(Paragraph(f"<b>Titel:</b> {titel}", heading_style))
        story.append(Paragraph(f"<b>Alarmierung:</b> {thema}", body_style))
        story.append(Spacer(1, 0.4*cm))

        # ABCDE-Tabelle
        if abcde_data:
            abcde = abcde_data.get('ABCDE', {})
            rows = [[Paragraph(f'<b>{k.upper()}=</b>', body_style),
                     Paragraph(str(v), body_style)]
                    for k, v in abcde.items() if str(v).strip()]
            if rows:
                tbl = Table(rows, colWidths=[1.8*cm, 15.2*cm])
                tbl.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), rl_font),
                    ('FONTSIZE', (0, 0), (-1, -1), fs),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.75, 0.75, 0.75)),
                    ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.93, 0.93, 0.93)),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(KeepTogether([
                    Paragraph('<b>ABCDE-Schema</b>', heading_style),
                    tbl,
                ]))
                story.append(Spacer(1, 0.3*cm))

        # Vitalwerte-Tabelle
        if vitalwerte:
            vw_map = {k: (lbl, unit) for k, lbl, unit in self._VITAL_LABELS}
            rows = []
            for k, v in vitalwerte.items():
                if str(v).strip():
                    lbl, unit = vw_map.get(k, (k, ''))
                    rows.append([Paragraph(f'<b>{lbl}</b>', body_style),
                                 Paragraph(f"{v} {unit}".strip(), body_style)])
            if rows:
                tbl = Table(rows, colWidths=[2.5*cm, 14.5*cm])
                tbl.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), rl_font),
                    ('FONTSIZE', (0, 0), (-1, -1), fs),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.Color(0.75, 0.75, 0.75)),
                    ('BACKGROUND', (0, 0), (0, -1), colors.Color(0.90, 0.95, 0.90)),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(KeepTogether([
                    Paragraph('<b>Vitalwerte / Messwerte</b>', heading_style),
                    tbl,
                ]))
                story.append(Spacer(1, 0.4*cm))

        # Inhalt
        story.append(Paragraph("<b>BERICHT:</b>", heading_style))
        
        # Inhalt in Paragraphen aufteilen
        for absatz in inhalt.split('\n'):
            if absatz.strip():
                # Prüfe ob es eine Überschrift ist (z.B. beginnt mit Nummer oder ist in Großbuchstaben)
                if absatz.strip().isupper() or absatz.strip()[0].isdigit():
                    story.append(Paragraph(absatz, heading_style))
                else:
                    story.append(Paragraph(absatz, body_style))
        
        # Reflexion
        if reflexion:
            story.append(Spacer(1, 0.5*cm))
            story.append(Paragraph("<b>EINSATZREFLEXION:</b>", heading_style))
            for absatz in reflexion.split('\n'):
                if absatz.strip():
                    story.append(Paragraph(absatz, body_style))
        
        # PDF bauen
        doc.build(story)
        
        return filepath
    
    def generate_word(self, titel: str, thema: str, inhalt: str, bericht_id: int,
                      font_family: str = "Arial", font_size: int = 11,
                      reflexion: str = "", abcde_data: dict = None, vitalwerte: dict = None) -> str:
        """
        Generiert ein Word-Dokument des Einsatzberichts
        
        Args:
            titel: Titel des Berichts
            thema: Thema des Berichts
            inhalt: Inhalt des Berichts
            bericht_id: ID des Berichts
        
        Returns:
            Pfad zur erstellten Word-Datei
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"bericht_{bericht_id}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)

        # Dokument erstellen
        doc = Document()

        # Standard-Schrift setzen (gilt für alle Absätze via Normal-Style)
        normal_style = doc.styles['Normal']
        normal_style.font.name = font_family or 'Arial'
        normal_style.font.size = Pt(font_size or 11)
        
        # Titel
        title = doc.add_heading('EINSATZBERICHT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Titel und Alarmierung
        p = doc.add_paragraph()
        p.add_run('Titel: ').bold = True
        p.add_run(titel)
        
        p = doc.add_paragraph()
        p.add_run('Alarmierung: ').bold = True
        p.add_run(thema)
        
        doc.add_paragraph()

        # ABCDE-Tabelle
        _fs = Pt(font_size or 11)
        if abcde_data:
            abcde = abcde_data.get('ABCDE', {})
            abcde_rows = [(k.upper() + '=', str(v)) for k, v in abcde.items() if str(v).strip()]
            if abcde_rows:
                h = doc.add_heading('ABCDE-Schema', 2)
                h.paragraph_format.keep_with_next = True
                tbl = doc.add_table(rows=len(abcde_rows), cols=2)
                tbl.style = 'Table Grid'
                for i, (key, val) in enumerate(abcde_rows):
                    tbl.cell(i, 0).text = key
                    tbl.cell(i, 1).text = val
                    for cell in [tbl.cell(i, 0), tbl.cell(i, 1)]:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = _fs
                doc.add_paragraph()

        # Vitalwerte-Tabelle
        if vitalwerte:
            vw_map = {k: (lbl, unit) for k, lbl, unit in self._VITAL_LABELS}
            vital_rows = []
            for k, v in vitalwerte.items():
                if str(v).strip():
                    lbl, unit = vw_map.get(k, (k, ''))
                    vital_rows.append((lbl, f"{v} {unit}".strip()))
            if vital_rows:
                h = doc.add_heading('Vitalwerte / Messwerte', 2)
                h.paragraph_format.keep_with_next = True
                tbl = doc.add_table(rows=len(vital_rows), cols=2)
                tbl.style = 'Table Grid'
                for i, (lbl, val) in enumerate(vital_rows):
                    tbl.cell(i, 0).text = lbl
                    tbl.cell(i, 1).text = val
                    for cell in [tbl.cell(i, 0), tbl.cell(i, 1)]:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = _fs
                doc.add_paragraph()
        
        # Bericht Überschrift
        doc.add_heading('BERICHT', 1)
        
        # Inhalt (keep_with_next auf Überschriften)
        for absatz in inhalt.split('\n'):
            if absatz.strip():
                # Prüfe ob es eine Überschrift ist
                if absatz.strip().isupper() or (absatz.strip() and absatz.strip()[0].isdigit()):
                    h = doc.add_heading(absatz.strip(), 2)
                    h.paragraph_format.keep_with_next = True
                else:
                    doc.add_paragraph(absatz)
        
        # Reflexion
        if reflexion:
            rh = doc.add_heading('EINSATZREFLEXION', 1)
            rh.paragraph_format.keep_with_next = True
            for absatz in reflexion.split('\n'):
                if absatz.strip():
                    doc.add_paragraph(absatz)
        
        # Speichern
        doc.save(filepath)
        
        return filepath
    
    def generate_both(self, titel: str, thema: str, inhalt: str, bericht_id: int) -> tuple:
        """
        Generiert sowohl PDF als auch Word-Dokument

        Returns:
            Tuple (pdf_path, word_path)
        """
        pdf_path = self.generate_pdf(titel, thema, inhalt, bericht_id)
        word_path = self.generate_word(titel, thema, inhalt, bericht_id)
        return pdf_path, word_path

    def generate_odf(self, titel: str, thema: str, inhalt: str, bericht_id: int,
                     font_family: str = "Arial", font_size: int = 11,
                     reflexion: str = "", abcde_data: dict = None, vitalwerte: dict = None) -> str:
        """Generiert ein ODF-Textdokument (.odt)"""
        from odf.opendocument import OpenDocumentText
        from odf.text import P, H
        from odf.style import Style, TextProperties, ParagraphProperties

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"bericht_{bericht_id}_{timestamp}.odt"
        filepath = os.path.join(self.output_dir, filename)

        doc = OpenDocumentText()

        # Standard-Textstil
        text_style = Style(name="CustomText", family="paragraph")
        text_style.addElement(TextProperties(
            fontfamily=font_family or "Arial",
            fontsize=f"{font_size or 11}pt"
        ))
        doc.styles.addElement(text_style)

        doc.text.addElement(H(outlinelevel=1, text="EINSATZBERICHT"))
        doc.text.addElement(P(text=""))
        doc.text.addElement(H(outlinelevel=2, text=f"Titel: {titel}"))
        doc.text.addElement(P(stylename=text_style, text=f"Alarmierung: {thema}"))
        doc.text.addElement(P(text=""))

        # ABCDE-Schema als Abschnitt
        if abcde_data:
            abcde = abcde_data.get('ABCDE', {})
            abcde_rows = [(k.upper() + '=', str(v)) for k, v in abcde.items() if str(v).strip()]
            if abcde_rows:
                doc.text.addElement(H(outlinelevel=2, text="ABCDE-Schema"))
                for key, val in abcde_rows:
                    doc.text.addElement(P(stylename=text_style, text=f"{key}  {val}"))
                doc.text.addElement(P(text=""))

        # Vitalwerte als Abschnitt
        if vitalwerte:
            vw_map = {k: (lbl, unit) for k, lbl, unit in self._VITAL_LABELS}
            vital_lines = []
            for k, v in vitalwerte.items():
                if str(v).strip():
                    lbl, unit = vw_map.get(k, (k, ''))
                    vital_lines.append(f"{lbl}:  {v} {unit}".strip())
            if vital_lines:
                doc.text.addElement(H(outlinelevel=2, text="Vitalwerte / Messwerte"))
                for line in vital_lines:
                    doc.text.addElement(P(stylename=text_style, text=line))
                doc.text.addElement(P(text=""))

        doc.text.addElement(H(outlinelevel=2, text="BERICHT:"))
        doc.text.addElement(P(text=""))

        for absatz in inhalt.split('\n'):
            if absatz.strip():
                if absatz.strip().isupper() or (absatz.strip() and absatz.strip()[0].isdigit()):
                    doc.text.addElement(H(outlinelevel=3, text=absatz))
                else:
                    doc.text.addElement(P(stylename=text_style, text=absatz))
            else:
                doc.text.addElement(P(text=""))

        if reflexion:
            doc.text.addElement(P(text=""))
            doc.text.addElement(H(outlinelevel=2, text="EINSATZREFLEXION:"))
            for absatz in reflexion.split('\n'):
                if absatz.strip():
                    doc.text.addElement(P(stylename=text_style, text=absatz))

        doc.save(filepath)
        return filepath

    def generate_pages(self, titel: str, thema: str, inhalt: str, bericht_id: int,
                       reflexion: str = "", abcde_data: dict = None, vitalwerte: dict = None) -> str:
        """Generiert ein Apple Pages-Dokument (.pages, iWork '09 Format)"""

        def escape(s: str) -> str:
            return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"bericht_{bericht_id}_{timestamp}.pages"
        filepath = os.path.join(self.output_dir, filename)

        abcde_lines = []
        if abcde_data:
            abcde = abcde_data.get('ABCDE', {})
            rows = [(k.upper() + '=', str(v)) for k, v in abcde.items() if str(v).strip()]
            if rows:
                abcde_lines = ["ABCDE-Schema:"] + [f"  {k}  {v}" for k, v in rows] + [""]
        vw_lines = []
        if vitalwerte:
            vw_map = {k: (lbl, unit) for k, lbl, unit in self._VITAL_LABELS}
            vw_parts = []
            for k, v in vitalwerte.items():
                if str(v).strip():
                    lbl, unit = vw_map.get(k, (k, ''))
                    vw_parts.append(f"  {lbl}:  {v} {unit}".strip())
            if vw_parts:
                vw_lines = ["Vitalwerte / Messwerte:"] + vw_parts + [""]
        reflexion_lines = (["EINSATZREFLEXION:", ""] + reflexion.split('\n')) if reflexion else []
        lines = (
            ["EINSATZBERICHT", "", f"Titel: {titel}", f"Alarmierung: {thema}", ""]
            + abcde_lines
            + vw_lines
            + ["BERICHT:", ""]
            + inhalt.split('\n')
            + ([""]+reflexion_lines if reflexion_lines else [])
        )

        para_xml = '\n'.join(
            f'        <sf:p sfa:ID="p{i}"><sf:content>{escape(ln)}</sf:content></sf:p>'
            for i, ln in enumerate(lines)
        )

        xml = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<sl:document xmlns:sl="http://developer.apple.com/namespaces/sl"\n'
            '  xmlns:sf="http://developer.apple.com/namespaces/sf"\n'
            '  xmlns:sfa="http://developer.apple.com/namespaces/sfa"\n'
            '  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"\n'
            '  sl:version="92082">\n'
            '  <sl:drawables>\n'
            '    <sl:section>\n'
            '      <sf:layout>\n'
            + para_xml + '\n'
            '      </sf:layout>\n'
            '    </sl:section>\n'
            '  </sl:drawables>\n'
            '</sl:document>'
        )

        with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("index.xml", xml.encode('utf-8'))

        return filepath

    # ------------------------------------------------------------------ #
    # Import-Methoden
    # ------------------------------------------------------------------ #

    def import_datei(self, filepath: str) -> dict:
        """Importiert einen Bericht aus einer Datei (docx / odt / pages)"""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.docx':
            return self.import_word(filepath)
        elif ext == '.odt':
            return self.import_odf(filepath)
        elif ext == '.pages':
            return self.import_pages(filepath)
        else:
            raise ValueError(f"Nicht unterstütztes Format: {ext}")

    def import_word(self, filepath: str) -> dict:
        """Importiert einen Bericht aus einer Word-Datei (.docx)"""
        from docx import Document
        doc = Document(filepath)
        titel = os.path.splitext(os.path.basename(filepath))[0]
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        return {'titel': titel, 'thema': '', 'inhalt': '\n'.join(lines)}

    def import_odf(self, filepath: str) -> dict:
        """Importiert einen Bericht aus einer ODF-Datei (.odt)"""
        from odf.opendocument import load
        from odf.text import P
        from odf import teletype
        doc = load(filepath)
        titel = os.path.splitext(os.path.basename(filepath))[0]
        lines = [teletype.extractText(p) for p in doc.text.getElementsByType(P)]
        return {'titel': titel, 'thema': '', 'inhalt': '\n'.join(lines)}

    def import_pages(self, filepath: str) -> dict:
        """Importiert einen Bericht aus einer Apple Pages-Datei (.pages)"""
        titel = os.path.splitext(os.path.basename(filepath))[0]
        lines = []
        try:
            with zipfile.ZipFile(filepath, 'r') as zf:
                names = zf.namelist()
                if 'index.xml' in names:
                    # Altes Pages-Format (iWork '09)
                    with zf.open('index.xml') as f:
                        content = f.read().decode('utf-8', errors='ignore')
                    try:
                        root = ET.fromstring(content)
                        for elem in root.iter():
                            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                            if tag == 'content' and elem.text and elem.text.strip():
                                lines.append(elem.text.strip())
                    except ET.ParseError:
                        lines = [t.strip() for t in re.findall(r'>([^<]+)<', content) if t.strip()]
                else:
                    # Neues Pages-Format: lesbaren Text extrahieren
                    for name in names:
                        if name.endswith('.txt'):
                            with zf.open(name) as f:
                                lines = f.read().decode('utf-8', errors='ignore').splitlines()
                            break
                    if not lines:
                        raise ValueError(
                            "Neue Apple Pages-Dateien (ab 2013) können nicht direkt importiert werden.\n"
                            "Bitte speichern Sie die Datei in Pages als .docx und importieren Sie dann die Word-Datei."
                        )
        except zipfile.BadZipFile:
            raise ValueError("Die .pages-Datei ist beschädigt oder kein unterstütztes Format.")
        return {'titel': titel, 'thema': '', 'inhalt': '\n'.join(lines)}
